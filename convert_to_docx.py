"""
Convert Markdown to Word Document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re


def markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document with formatting."""

    # Read markdown content
    with open(md_file, "r", encoding="utf-8") as f:
        content = f.read()

    # Create Word document
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Process content line by line
    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                if code_content:
                    p = doc.add_paragraph()
                    p.style = "No Spacing"
                    for code_line in code_content:
                        run = p.add_run(code_line + "\n")
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Handle tables
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_rows = []

            # Skip separator rows
            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                i += 1
                continue

            # Parse table row
            cells = [cell.strip() for cell in line.split("|")[1:-1]]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table, create it
            if table_rows:
                num_cols = max(len(row) for row in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = "Table Grid"

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            # Clean markdown formatting
                            clean_text = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_data)
                            clean_text = re.sub(r"\*([^*]+)\*", r"\1", clean_text)
                            cell.text = clean_text

                doc.add_paragraph()  # Space after table
            table_rows = []
            in_table = False

        # Handle headers
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        elif line.startswith("#### "):
            p = doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # Handle horizontal rules
        if line.strip() == "---":
            doc.add_paragraph("─" * 50)
            i += 1
            continue

        # Handle bullet points
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, text)
            i += 1
            continue

        # Handle numbered lists
        match = re.match(r"^(\d+)\.\s+(.+)$", line.strip())
        if match:
            text = match.group(2)
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, text)
            i += 1
            continue

        # Handle empty lines
        if line.strip() == "":
            i += 1
            continue

        # Handle regular paragraphs
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        i += 1

    # Handle any remaining table
    if in_table and table_rows:
        num_cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = "Table Grid"

        for row_idx, row_data in enumerate(table_rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    clean_text = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_data)
                    clean_text = re.sub(r"\*([^*]+)\*", r"\1", clean_text)
                    cell.text = clean_text

    # Save document
    doc.save(docx_file)
    print(f"✓ Successfully converted to: {docx_file}")


def add_formatted_text(paragraph, text):
    """Add text with bold and italic formatting."""
    # Pattern to find **bold** and *italic* text
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\$[^$]+\$|\$\$[^$]+\$\$|[^*`$]+)"
    parts = re.findall(pattern, text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif part.startswith("$$") and part.endswith("$$"):
            run = paragraph.add_run(part[2:-2])
            run.italic = True
        elif part.startswith("$") and part.endswith("$"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


if __name__ == "__main__":
    import os

    base_path = r"C:\Users\nduta\OneDrive\Desktop\Projects\Resource-Allocation-Constrained-Optimization"
    md_file = os.path.join(base_path, "assignment_answers.md")
    docx_file = os.path.join(base_path, "assignment_answers.docx")

    markdown_to_docx(md_file, docx_file)
